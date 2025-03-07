import os
import shutil
from pathlib import Path
from typing import List, Optional, Dict, Any, Tuple
from fastapi import UploadFile, HTTPException
from sqlalchemy.orm import Session
from datetime import datetime

from ..models.models import Document
from ..models.schemas import DocumentCreate, DocumentResponse
from ..core.config import DOCUMENT_DIR
from . import ai_service

async def save_upload_file(file: UploadFile, destination: Path) -> Path:
    """Save an uploaded file to the specified destination."""
    # Create destination directory if it doesn't exist
    os.makedirs(destination, exist_ok=True)
    
    # Generate a unique filename with timestamp
    timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
    filename = f"{timestamp}_{file.filename}"
    file_path = Path(destination) / filename
    
    # Write file
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    return file_path

async def create_document(
    db: Session, 
    file: UploadFile, 
    document_data: DocumentCreate
) -> DocumentResponse:
    """Create a new document record and save the uploaded file."""
    # Save the file
    file_path = await save_upload_file(file, DOCUMENT_DIR)
    
    # Extract content from the file based on file type
    content = ""
    try:
        file_extension = os.path.splitext(file.filename)[1].lower() if file.filename else ""
        
        if file_extension in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json']:
            # Text files
            with open(file_path, "r", errors="ignore") as f:
                content = f.read()
        elif file_extension in ['.pdf']:
            # PDF files
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    content = ""
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        content += page.extract_text() + "\n\n"
            except ImportError:
                content = "[PDF content extraction not available - PyPDF2 not installed]"
        elif file_extension == '.doc':
            # Word .doc files
            try:
                # Try using textract
                import textract
                content = textract.process(file_path).decode("utf-8", errors="ignore")
            except ImportError:
                try:
                    # Alternative: try using antiword
                    import subprocess
                    result = subprocess.run(['antiword', file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    if result.returncode == 0:
                        content = result.stdout.decode("utf-8", errors="ignore")
                    else:
                        # Fall back to docx2txt
                        import docx2txt
                        content = docx2txt.process(file_path)
                except Exception as e:
                    content = f"[DOC content extraction error: {str(e)}]"
        elif file_extension == '.docx':
            # Word .docx files
            try:
                import docx
                doc = docx.Document(file_path)
                content = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                try:
                    # Alternative: try using docx2txt
                    import docx2txt
                    content = docx2txt.process(file_path)
                except ImportError:
                    content = "[DOCX content extraction not available - python-docx not installed]"
        else:
            content = f"[Content extraction not supported for {file_extension} files]"
    except Exception as e:
        print(f"Error extracting content: {e}")
        content = f"[Error extracting content: {str(e)}]"
    
    # Create database record
    db_document = Document(
        title=document_data.title,
        file_path=str(file_path),
        content=content,
        category=document_data.category,
        offline_available=document_data.offline_available
    )
    
    db.add(db_document)
    db.commit()
    db.refresh(db_document)
    
    # Add to vector store for RAG as a single vector, not chunked
    try:
        # Store document metadata with the full text content
        metadata = {
            "title": document_data.title,
            "category": document_data.category,
            "document_id": db_document.id,
            "source": "document"
        }
        
        # Add to Pinecone as a single document
        vector_id = ai_service.add_document_to_vector_store(db_document.id, content, metadata)
        
        # Update the document with the vector ID
        db_document.vector_id = vector_id
        db.commit()
        db.refresh(db_document)
        
        print(f"Document added to vector store with ID: {vector_id}")
    except Exception as e:
        print(f"Error adding document to vector store: {e}")
    
    return db_document

def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for better semantic search."""
    if not text:
        return []
        
    # Input validation
    if not isinstance(text, str):
        print(f"Warning: chunk_text received non-string input: {type(text)}")
        try:
            text = str(text)
        except:
            return []
            
    if chunk_size <= 0:
        chunk_size = 1000
    if overlap < 0 or overlap >= chunk_size:
        overlap = min(200, chunk_size // 5)
        
    try:
        # Split text into sentences (simple approach)
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for sentence in sentences:
            sentence_size = len(sentence)
            
            # If adding this sentence would exceed the chunk size, finalize the chunk
            if current_size + sentence_size > chunk_size and current_chunk:
                chunks.append(" ".join(current_chunk))
                
                # Keep some sentences for overlap
                overlap_sentences = []
                overlap_size = 0
                
                # Work backwards through current_chunk to find sentences for overlap
                for s in reversed(current_chunk):
                    if overlap_size + len(s) <= overlap:
                        overlap_sentences.insert(0, s)
                        overlap_size += len(s)
                    else:
                        break
                
                # Start a new chunk with overlap sentences
                current_chunk = overlap_sentences
                current_size = overlap_size
            
            current_chunk.append(sentence)
            current_size += sentence_size
        
        # Add the last chunk if it's not empty
        if current_chunk:
            chunks.append(" ".join(current_chunk))
        
        # If no chunks were created (perhaps the regex didn't split anything),
        # fall back to basic chunking
        if not chunks and text:
            # Simple fallback: chunk by character
            for i in range(0, len(text), chunk_size - overlap):
                end = min(i + chunk_size, len(text))
                chunks.append(text[i:end])
                
                # If we've reached the end of the text, stop
                if end == len(text):
                    break
        
        return chunks
    except Exception as e:
        print(f"Error in chunk_text: {e}")
        # Emergency fallback
        return [text[:chunk_size]] if text else []

def get_document(db: Session, document_id: int) -> Optional[Document]:
    """Get a document by ID."""
    return db.query(Document).filter(Document.id == document_id).first()

def get_documents(
    db: Session, 
    skip: int = 0, 
    limit: int = 100, 
    category: Optional[str] = None
) -> List[Document]:
    """Get all documents, optionally filtered by category."""
    query = db.query(Document)
    
    if category:
        query = query.filter(Document.category == category)
    
    return query.offset(skip).limit(limit).all()

def update_document(
    db: Session, 
    document_id: int, 
    document_data: DocumentCreate
) -> Optional[Document]:
    """Update a document's metadata."""
    db_document = get_document(db, document_id)
    
    if not db_document:
        return None
    
    # Update fields
    for key, value in document_data.dict(exclude_unset=True).items():
        setattr(db_document, key, value)
    
    db.commit()
    db.refresh(db_document)
    
    return db_document

def delete_document(db: Session, document_id: int) -> bool:
    """Delete a document and its file."""
    db_document = get_document(db, document_id)
    
    if not db_document:
        return False
    
    # Delete the file
    file_path = Path(db_document.file_path)
    if file_path.exists():
        file_path.unlink()
    
    # Delete from database
    db.delete(db_document)
    db.commit()
    
    return True

def search_documents(query: str, top_k: int = 3) -> List[Dict[str, Any]]:
    """Search for documents using RAG."""
    try:
        results = ai_service.search_similar_documents(query, top_k)
        
        # Debug the results
        print(f"Found {len(results)} similar documents for query: '{query}'")
        
        # Ensure each result has the necessary metadata
        for i, result in enumerate(results):
            # Check if it has metadata
            metadata = None
            if hasattr(result, 'metadata'):
                metadata = result.metadata
            elif isinstance(result, dict) and 'metadata' in result:
                metadata = result['metadata']
                
            if metadata and not metadata.get('text'):
                # Try to retrieve the document and add its content
                doc_id = metadata.get('document_id')
                if doc_id:
                    try:
                        from sqlalchemy.orm import Session
                        from ..db.database import SessionLocal
                        from ..models.models import Document
                        
                        db = SessionLocal()
                        doc = db.query(Document).filter(Document.id == doc_id).first()
                        if doc:
                            metadata['text'] = doc.content
                            print(f"Added content for document {doc_id} to result {i}")
                        db.close()
                    except Exception as e:
                        print(f"Error retrieving document content: {e}")
        
        return results
    except Exception as e:
        print(f"Error searching documents: {e}")
        return []

def extract_relevant_sections(query: str, docs: List[Dict[str, Any]], max_tokens: int = 8000) -> Tuple[str, List[str]]:
    """
    Extract only the most relevant sections from the retrieved documents 
    to avoid context length issues.
    
    Args:
        query: The user's query
        docs: The list of document matches from Pinecone
        max_tokens: Maximum tokens to include in the context (approximate)
        
    Returns:
        tuple containing: (context_text, source_titles)
    """
    if not docs:
        return "", []
    
    # Initialize
    context = "Relevant information from our knowledge base:\n\n"
    sources = []
    total_chars = len(context)  # approximate token count by chars (rough estimate: 1 token = 4 chars)
    max_chars = max_tokens * 4  # rough conversion
    
    # Track processed documents to avoid duplicates
    processed_doc_ids = set()
    
    # First, extract all document contents and split them into semantic sections
    document_sections = []
    
    for doc in docs:
        # Extract metadata based on the response format
        doc_text = ""
        doc_title = ""
        doc_id = None
        
        if hasattr(doc, 'metadata') and doc.metadata:
            metadata = doc.metadata
            doc_title = metadata.get('title', "")
            doc_id = metadata.get('document_id')
            doc_text = metadata.get('text', '')
        elif isinstance(doc, dict):
            metadata = doc.get('metadata', {})
            doc_title = metadata.get('title', "")
            doc_id = metadata.get('document_id')
            doc_text = metadata.get('text', '')
        
        # Skip if empty or already processed
        if not doc_text or not doc_id or doc_id in processed_doc_ids:
            continue
            
        processed_doc_ids.add(doc_id)
        
        # Add to sources
        if doc_title and doc_title not in sources:
            sources.append(doc_title)
        
        # Split the document into semantic sections
        sections = split_document_into_sections(doc_text)
        
        # Add sections with metadata
        for section_title, section_content in sections.items():
            document_sections.append({
                "doc_id": doc_id,
                "doc_title": doc_title,
                "section_title": section_title,
                "content": section_content,
                "chars": len(section_content)
            })
    
    # Sort sections by length (shortest first to maximize variety)
    document_sections.sort(key=lambda x: x["chars"])
    
    # Score each section by relevance to the query
    scored_sections = []
    for section in document_sections:
        # Simple relevance scoring: count query terms in section
        relevance_score = calculate_relevance(query, section["content"])
        scored_sections.append({
            **section,
            "relevance": relevance_score
        })
    
    # Sort by relevance score (highest first)
    scored_sections.sort(key=lambda x: x["relevance"], reverse=True)
    
    # Add sections to context until we reach the token limit
    for section in scored_sections:
        # Calculate potential size with this section
        section_text = f"--- {section['doc_title']} ({section['section_title']}) ---\n{section['content']}\n\n"
        potential_size = total_chars + len(section_text)
        
        # Check if adding this section would exceed our limit
        if potential_size > max_chars:
            # If this is the first section, take a portion of it
            if total_chars == len(context):
                # Take as much as we can
                chars_available = max_chars - total_chars
                truncated_content = section['content'][:chars_available] + "..."
                section_text = f"--- {section['doc_title']} ({section['section_title']}) ---\n{truncated_content}\n\n"
                context += section_text
            break
            
        # Add the section to our context
        context += section_text
        total_chars = potential_size
    
    return context, sources

def calculate_relevance(query: str, text: str) -> float:
    """
    Calculate the relevance of a text section to the query.
    Higher score = more relevant.
    """
    # Simple implementation using term frequency
    query_terms = set(query.lower().split())
    text_lower = text.lower()
    
    # Remove common stopwords from the query terms
    stopwords = {"the", "a", "an", "in", "on", "at", "to", "for", "with", "by", "about", "like", "through", "over", "of"}
    query_terms = query_terms - stopwords
    
    if not query_terms:
        return 0
    
    score = 0
    for term in query_terms:
        if term in text_lower:
            # Add points based on frequency
            freq = text_lower.count(term)
            score += min(freq * 10, 100)  # Cap the score per term
            
            # Bonus for exact phrase matches
            if query.lower() in text_lower:
                score += 200
    
    # Normalize by length to favor concise sections
    length_factor = max(1, min(len(text) / 500, 3))  # 1 for short texts, up to 3 for very long texts
    
    # Final score: relevance divided by length factor
    return score / length_factor

def split_document_into_sections(text: str, max_section_chars: int = 2000) -> Dict[str, str]:
    """
    Split a document into semantic sections for better retrieval.
    
    Returns a dictionary of {section_title: section_content}
    """
    sections = {}
    
    # Try to split by double newlines first (paragraphs)
    paragraphs = text.split("\n\n")
    
    # If we have very few paragraphs, try single newlines
    if len(paragraphs) <= 3:
        paragraphs = text.split("\n")
    
    # Group paragraphs into sections
    current_section = []
    current_section_title = "Section 1"
    section_counter = 1
    current_length = 0
    
    for i, para in enumerate(paragraphs):
        para_stripped = para.strip()
        
        # Skip empty paragraphs
        if not para_stripped:
            continue
            
        # Check if this looks like a header (short, ends with colon, or all caps)
        is_header = (len(para_stripped) < 50 and 
                   (para_stripped.endswith(':') or para_stripped.isupper() or
                    all(c.isupper() or not c.isalpha() for c in para_stripped[:10])))
        
        if is_header:
            # Save the previous section if it exists
            if current_section:
                sections[current_section_title] = "\n\n".join(current_section)
                
            # Start a new section with this header as title
            current_section_title = para_stripped
            current_section = []
            current_length = 0
        else:
            # Check if adding this paragraph would make section too long
            if current_length + len(para_stripped) > max_section_chars and current_section:
                # Save current section and start a new one
                sections[current_section_title] = "\n\n".join(current_section)
                section_counter += 1
                current_section_title = f"Section {section_counter}"
                current_section = [para_stripped]
                current_length = len(para_stripped)
            else:
                # Add to current section
                current_section.append(para_stripped)
                current_length += len(para_stripped)
    
    # Add the last section
    if current_section:
        sections[current_section_title] = "\n\n".join(current_section)
    
    # If we couldn't find good sections, return the whole text as one section
    if not sections:
        sections["Full Text"] = text
        
    return sections 